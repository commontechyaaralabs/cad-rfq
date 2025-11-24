import base64
import io
import json
import logging
import mimetypes
import os
import re
import subprocess
from typing import List, Tuple, Union, Optional, Dict

import cv2
import numpy as np
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from typing import List as TypingList
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from google import genai  # type: ignore[reportMissingImports]
from google.genai import types  # type: ignore[reportMissingImports]
from docx import Document  # type: ignore[reportMissingImports]

# ----------------- Config -----------------

def _resolve_project_id() -> str:
    """Determine the active GCP project ID from env vars or gcloud CLI."""
    # Prefer explicit env configuration
    env_keys = ("GOOGLE_CLOUD_PROJECT", "GCLOUD_PROJECT", "PROJECT_ID")
    for key in env_keys:
        value = os.environ.get(key)
        if value:
            return value

    # Fallback to `gcloud config get-value project`
    try:
        result = subprocess.check_output(
            ["gcloud", "config", "get-value", "project"],
            stderr=subprocess.STDOUT,
            text=True,
        ).strip()
        if result and result.lower() != "(unset)":
            return result
    except (subprocess.CalledProcessError, FileNotFoundError, OSError):
        pass

    raise RuntimeError(
        "Unable to determine GCP project ID. "
        "Set GOOGLE_CLOUD_PROJECT (or run `gcloud config set project <id>`)."
    )


DEFAULT_PROJECT = "playgroundai-470111"
_bootstrap_logger = logging.getLogger(__name__)


def _resolve_project_id() -> str:
    """Determine the active GCP project ID from env vars, gcloud config, or a fallback constant."""
    env_keys = ("GOOGLE_CLOUD_PROJECT", "GCLOUD_PROJECT", "PROJECT_ID")
    for key in env_keys:
        value = os.environ.get(key)
        if value:
            _bootstrap_logger.info("[CONFIG] Using project ID from %s", key)
            return value

    try:
        result = subprocess.check_output(
            ["gcloud", "config", "get-value", "project"],
            stderr=subprocess.STDOUT,
            text=True,
        ).strip()
        if result and result.lower() != "(unset)":
            _bootstrap_logger.info("[CONFIG] Using project ID from gcloud config")
            return result
    except (subprocess.CalledProcessError, FileNotFoundError, OSError):
        _bootstrap_logger.info("[CONFIG] gcloud project lookup failed; falling back to default")

    _bootstrap_logger.warning(
        "[CONFIG] Falling back to default project ID '%s'. "
        "Set GOOGLE_CLOUD_PROJECT to override.",
        DEFAULT_PROJECT,
    )
    return DEFAULT_PROJECT


PROJECT = _resolve_project_id()          # active GCP project ID
REGION = "us-east4"                      # Vertex AI region
MODEL = "gemini-2.5-pro"                 # Gemini model name

# ----------------- Logging Setup -----------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("inspection.log", mode="a", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger(__name__)

# ----------------- FastAPI App -----------------
app = FastAPI()

# Enable CORS for frontend
# Get allowed origins from environment or use defaults
allowed_origins_env = os.getenv("CORS_ORIGINS", "")
if allowed_origins_env:
    # Split comma-separated origins from environment
    allowed_origins = [origin.strip() for origin in allowed_origins_env.split(",")]
else:
    # Default origins for local development
    allowed_origins = [
        "http://localhost:3000",  # Next.js default port
        "http://localhost:3001",  # Alternative Next.js port
        "http://127.0.0.1:3000",
        "http://127.0.0.1:3001",
    ]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ----------------- Comparison Prompt Strategies -----------------


class ComparisonPromptStrategy:
    """Base strategy for building RFQ vs CAD comparison prompts."""

    part_key: str = "spark_plug"
    aliases: Tuple[str, ...] = ()

    def build_prompt(self) -> str:
        raise NotImplementedError


class SparkPlugComparisonStrategy(ComparisonPromptStrategy):
    part_key = "spark_plug"
    aliases = ("spark plug", "sparkplug")

    _PROMPT = (
        "You are an expert sourcing engineer specializing in automotive components. "
        "Compare the attached RFQ document, which requests a standard OEM part, with the attached engineering drawing. "
        "Your task is to verify if the drawing dimensionally corresponds to the specific part number mentioned in the RFQ.\n\n"
        "IMPORTANT CONTEXT: This RFQ is for a specific, off-the-shelf product ('Bosch FR7DC'). "
        "Therefore, the drawing should be treated as a dimensional reference to confirm the part's form and fit. "
        "Requirements like material composition, manufacturer branding, heat range, or performance specifications "
        "are defined by the part number itself and are not expected to be repeated on the reference drawing.\n\n"
        "INSTRUCTIONS:\n"
        "1. From the RFQ, identify the required part number and extract all key dimensional requirements "
        "(e.g., thread size, overall length, hex size, electrode gap).\n"
        "2. Examine the CAD drawing and list the corresponding dimensions it provides.\n"
        "3. Compare the dimensional requirements from the RFQ against the dimensions found in the drawing.\n"
        "4. Treat each physical dimension as a SINGLE metric. If RFQ and CAD use different labels for the same thing "
        "(e.g., 'Insulator Nose Length' vs 'Insulator Length', 'Shell Diameter' vs 'Maximum Shell Diameter', "
        "'Electrode Gap' vs 'Electrode G'), you MUST merge them into one combined metric instead of creating "
        "separate entries. Always prefer the RFQ label as the primary name and mention the CAD label in parentheses.\n"
        "   - Example merged metric in rfq_requirements: "
        "     'Insulator Nose Length (CAD: Insulator Length): 42 mm'.\n"
        "   - The same metric in cad_findings should use the same merged name and the value from the drawing.\n"
        "5. Do NOT create duplicate metrics that only differ by wording. There must be at most one RFQ/CAD pair "
        "for each actual dimension.\n"
        "6. If all specified dimensions in the RFQ are correctly represented in the drawing, mark 'match': true. "
        "The absence of non-dimensional information (like material or brand) on the drawing does not constitute a mismatch. "
        "Mark 'match': false only if there is a direct conflict in a dimension.\n"
        "7. If a mismatch is found, clearly identify the specific dimensional discrepancy. Otherwise, this section can be empty.\n"
        "8. End with a short recommendation for the procurement team based on your findings.\n\n"
        "OUTPUT STRICTLY AS VALID JSON USING THIS STRUCTURE:\n"
        "{\n"
        '  "match": true or false,\n'
        '  "confidence": "High|Medium|Low",\n'
        '  "summary": "Overall assessment in one sentence",\n'
        '  "rfq_requirements": ["Requirement 1", ...],\n'
        '  "cad_findings": ["Finding 1", ...],\n'
        '  "mismatches": ["Mismatch detail", ...],\n'
        '  "recommendations": "Clear next actions"\n'
        "}\n"
    )

    def build_prompt(self) -> str:
        return self._PROMPT


class BrakeDiscComparisonStrategy(ComparisonPromptStrategy):
    part_key = "brake_disc"
    aliases = ("brake disc", "brakedisc", "brake disk", "brakedisk")

    _PROMPT = (
        "You are an expert automotive and mechanical engineer specializing in brake disc (rotor) design. "
        "You will be given:\n"
        "1) An RFQ document describing a specific front brake disc.\n"
        "2) An engineering drawing (CAD) of a brake disc.\n\n"
        "Your job is to determine whether the CAD drawing dimensionally matches the RFQ requirements for the same part.\n\n"
        "GENERAL RULES:\n"
        "- Treat the RFQ as the source of truth for required dimensions.\n"
        "- Focus ONLY on dimensional and geometric information: diameters, thicknesses, heights, PCD, bore sizes, hole count/diameter, etc.\n"
        "- Do NOT treat missing non-dimensional data (material grade, surface finish, coating, packaging, quantity, delivery terms, etc.) as a mismatch.\n"
        "- Do NOT invent or guess dimensions that are not visible in the drawing.\n\n"
        "STEP 1 – Extract RFQ Requirements:\n"
        "1. Carefully read the RFQ and extract EVERY dimensional requirement for the brake disc, including (as applicable):\n"
        "   - Overall outer diameter\n"
        "   - Effective braking diameter / friction ring diameter\n"
        "   - Overall thickness of the disc\n"
        "   - Individual thicknesses or heights (hat height, hub heights, top/bottom sections, etc.)\n"
        "   - Center bore / hub bore diameter\n"
        "   - Any intermediate diameters (steps, shoulders, register diameters)\n"
        "   - Bolt circle / PCD, number of bolt holes, and bolt hole diameters (including multiple hole groups)\n"
        "   - Any dimensional tolerances or limits that are explicitly given.\n"
        "2. For each metric, create a single entry in the form 'Metric Name: value + unit'. "
        "   Always include the unit (for example '300 mm', 'Ø195 mm').\n"
        "3. Metric names must be clear and human-readable, for example: "
        "'Outer Diameter', 'PCD (Bolt Circle Diameter)', 'Center Bore', 'Overall Thickness', "
        "'Hat Height', 'Hub Height (Top Section)', 'Hub Height (Bottom Section)', "
        "'Bolt Hole Diameter (Group 1)', 'Bolt Hole Diameter (Group 2)'.\n"
        "4. Ignore commercial sections such as quantity, pricing, delivery, payment terms, certifications, etc.\n\n"
        "STEP 2 – Read the CAD Drawing:\n"
        "1. Examine all views and sections (e.g., front view, top view, sections A–A, B–B, etc.).\n"
        "2. For each RFQ metric, try to find the exact corresponding dimension on the drawing.\n"
        "   - If the drawing uses a different label for the same concept, still treat it as the same metric. "
        "     Examples: 'Overall Height' vs 'Hat Height', 'Center Bore' vs 'Hub Bore', "
        "'Effective Braking Diameter' vs 'Ø195 braking surface'.\n"
        "3. For each metric that you can read from the drawing, create an entry using the SAME metric name you used in the RFQ list, "
        "in the form 'Metric Name: value + unit'.\n"
        "4. If a metric from the RFQ cannot be found on the drawing, still create an entry using the SAME metric name but with the value '—'. "
        "   Example: 'Minimum Thickness: —'.\n\n"
        "STEP 3 – Compare RFQ vs CAD:\n"
        "1. Compare every RFQ metric with the corresponding CAD metric.\n"
        "2. If the CAD dimension is clearly equal or equivalent (same value, or obviously the same within normal rounding such as 80 mm vs 80.0 mm), "
        "treat it as a match.\n"
        "3. If the CAD dimension is different, treat it as a mismatch and describe the difference explicitly.\n"
        "4. If the CAD dimension is missing (value '—'), treat that as a potential risk and record it in the mismatches list as 'missing dimension on drawing'.\n"
        "5. Set 'match' to true ONLY IF all RFQ dimensional metrics are satisfied by the drawing (no conflicting values and no critical missing dimensions). "
        "If any dimension is conflicting or clearly missing, set 'match' to false.\n\n"
        "STEP 4 – Recommendations:\n"
        "1. Based on the comparison, give clear next actions for the procurement/engineering team, such as:\n"
        "   - 'Proceed – drawing is fully consistent with RFQ dimensions.'\n"
        "   - 'Request updated drawing with corrected PCD.'\n"
        "   - 'Ask supplier to clarify missing minimum thickness dimension.'\n\n"
        "OUTPUT FORMAT (CRITICAL):\n"
        "- You MUST output ONLY valid JSON.\n"
        "- Do NOT wrap the JSON in markdown code fences.\n"
        "- Do NOT include any prose before or after the JSON.\n"
        "- Arrays must contain plain strings; do not nest objects inside them.\n"
        "- Each RFQ and CAD item must be of the form 'Metric Name: value'.\n\n"
        "OUTPUT STRICTLY AS VALID JSON USING THIS STRUCTURE:\n"
        "{\n"
        '  "match": true or false,\n'
        '  "confidence": "High" or "Medium" or "Low",\n'
        '  "summary": "Overall assessment in one clear sentence",\n'
        '  "rfq_requirements": ["Metric 1: value", "Metric 2: value", ...],\n'
        '  "cad_findings": ["Metric 1: value", "Metric 2: value", ...],\n'
        '  "mismatches": ["Metric: RFQ value vs CAD value – explanation", ...],\n'
        '  "recommendations": "Concrete next actions for the team"\n'
        "}\n"
    )

    def build_prompt(self) -> str:
        return self._PROMPT


class HornComparisonStrategy(ComparisonPromptStrategy):
    part_key = "horn"
    aliases = ("horn", "horn assembly", "automotive horn", "car horn")

    _PROMPT = (
        "You are an expert sourcing engineer specializing in automotive horn assemblies. "
        "You will be given:\n"
        "1) An RFQ document describing a specific horn or horn assembly.\n"
        "2) An engineering drawing (CAD) of a horn or horn assembly.\n\n"
        "Your job is to determine whether the CAD drawing dimensionally matches the RFQ requirements for the same part.\n\n"
        "GENERAL RULES:\n"
        "- Treat the RFQ as the source of truth for required dimensions.\n"
        "- Focus ONLY on dimensional and geometric information: diameters, depths, widths, mounting bracket dimensions, "
        "bolt circle/PCD, hole diameters, heights, connector locations, etc.\n"
        "- Do NOT treat missing non-dimensional data (material grade, surface finish, coating, packaging, quantity, delivery terms, etc.) as a mismatch.\n"
        "- Do NOT invent or guess dimensions that are not visible in the drawing.\n\n"
        "STEP 1 – Extract RFQ Requirements:\n"
        "1. Carefully read the RFQ and extract EVERY dimensional requirement for the horn, including (as applicable):\n"
        "   - Overall diameter or envelope size\n"
        "   - Horn body depth/width\n"
        "   - Mounting bracket length, thickness, and hole diameter(s)\n"
        "   - Bolt circle / PCD, number and position of mounting holes\n"
        "   - Center height, offset height, or mounting height from reference\n"
        "   - Connector type/location if dimensionally specified\n"
        "   - Any dimensional tolerances or limits that are explicitly given.\n"
        "2. For each metric, create a single entry in the form 'Metric Name: value + unit'. "
        "   Always include the unit (for example '120 mm', 'Ø85 mm').\n"
        "3. Metric names must be clear and human-readable, for example: "
        "'Overall Diameter', 'Horn Body Depth', 'Mounting Bracket Length', 'Bolt Circle (PCD)', "
        "'Mounting Hole Diameter', 'Center Height', 'Connector Location'.\n"
        "4. Ignore commercial sections such as quantity, pricing, delivery, payment terms, certifications, etc.\n\n"
        "STEP 2 – Read the CAD Drawing:\n"
        "1. Examine all views and sections (e.g., front view, top view, side view, sections, etc.).\n"
        "2. For each RFQ metric, try to find the exact corresponding dimension on the drawing.\n"
        "   - If the drawing uses a different label for the same concept, still treat it as the same metric. "
        "     Examples: 'Overall Diameter' vs 'Envelope Diameter', 'Mounting Height' vs 'Center Height', "
        "'Bolt Circle' vs 'PCD'.\n"
        "3. For each metric that you can read from the drawing, create an entry using the SAME metric name you used in the RFQ list, "
        "in the form 'Metric Name: value + unit'.\n"
        "4. If a metric from the RFQ cannot be found on the drawing, still create an entry using the SAME metric name but with the value '—'. "
        "   Example: 'Mounting Bracket Thickness: —'.\n\n"
        "STEP 3 – Compare RFQ vs CAD:\n"
        "1. Compare every RFQ metric with the corresponding CAD metric.\n"
        "2. If the CAD dimension is clearly equal or equivalent (same value, or obviously the same within normal rounding such as 120 mm vs 120.0 mm), "
        "treat it as a match.\n"
        "3. If the CAD dimension is different, treat it as a mismatch and describe the difference explicitly.\n"
        "4. If the CAD dimension is missing (value '—'), treat that as a potential risk and record it in the mismatches list as 'missing dimension on drawing'.\n"
        "5. Set 'match' to true ONLY IF all RFQ dimensional metrics are satisfied by the drawing (no conflicting values and no critical missing dimensions). "
        "If any dimension is conflicting or clearly missing, set 'match' to false.\n\n"
        "STEP 4 – Recommendations:\n"
        "1. Based on the comparison, give clear next actions for the procurement/engineering team, such as:\n"
        "   - 'Proceed – drawing is fully consistent with RFQ dimensions.'\n"
        "   - 'Request updated drawing with corrected bolt circle diameter.'\n"
        "   - 'Ask supplier to clarify missing mounting bracket thickness dimension.'\n\n"
        "OUTPUT FORMAT (CRITICAL):\n"
        "- You MUST output ONLY valid JSON.\n"
        "- Do NOT wrap the JSON in markdown code fences.\n"
        "- Do NOT include any prose before or after the JSON.\n"
        "- Arrays must contain plain strings; do not nest objects inside them.\n"
        "- Each RFQ and CAD item must be of the form 'Metric Name: value'.\n\n"
        "OUTPUT STRICTLY AS VALID JSON USING THIS STRUCTURE:\n"
        "{\n"
        '  "match": true or false,\n'
        '  "confidence": "High" or "Medium" or "Low",\n'
        '  "summary": "Overall assessment in one clear sentence",\n'
        '  "rfq_requirements": ["Metric 1: value", "Metric 2: value", ...],\n'
        '  "cad_findings": ["Metric 1: value", "Metric 2: value", ...],\n'
        '  "mismatches": ["Metric: RFQ value vs CAD value – explanation", ...],\n'
        '  "recommendations": "Concrete next actions for the team"\n'
        "}\n"
    )

    def build_prompt(self) -> str:
        return self._PROMPT


# ----------------- Gemini Client -----------------
class GeminiClient:
    """Wrapper around Gemini 2.5 Pro via Vertex AI SDK."""

    def __init__(self, project: str, region: str, model_name: str):
        self.client = genai.Client(
            vertexai=True,
            project=project,
            location=region,
        )
        self.model_name = model_name

    def _create_file_part(self, file_bytes: bytes, mime_type: Optional[str]) -> types.Part:
        file_type = "PDF" if mime_type == "application/pdf" else "image"
        logger.info(f"Sending {file_type} to Gemini model {self.model_name}")

        file_b64 = base64.b64encode(file_bytes).decode("utf-8")

        return types.Part(
            inline_data=types.Blob(
                mime_type=mime_type or "application/octet-stream",
                data=file_b64
            )
        )

    def chat(self, file_bytes: bytes, mime_type: str, prompt: str) -> str:
        """Send single file + prompt to Gemini and return text response."""
        return self.chat_with_files(prompt, [(file_bytes, mime_type)])

    def chat_with_files(self, prompt: str, files: List[Tuple[Union[bytes, str], Optional[str]]]) -> str:
        """Send multiple files + prompt to Gemini and return text response."""
        contents: List[Union[str, types.Part]] = [prompt]
        for data, mime in files:
            if isinstance(data, bytes):
                contents.append(self._create_file_part(data, mime))
            else:
                contents.append(str(data))

        response = self.client.models.generate_content(
            model=self.model_name,
            contents=contents,
            config={
                "max_output_tokens": 16384,
                "temperature": 0,  # Set to 0 for maximum determinism
            },
        )

        if hasattr(response, 'candidates') and response.candidates:
            candidate = response.candidates[0]
            if hasattr(candidate, 'finish_reason'):
                if candidate.finish_reason == 'MAX_TOKENS':
                    logger.warning("⚠ Response may have been truncated due to token limit")
                logger.info(f"Finish reason: {candidate.finish_reason}")

        return response.text


# ----------------- Welding Inspector -----------------
class WeldingInspector:
    def __init__(self, client: GeminiClient):
        self.client = client
        self.word_mime_types = {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }
        strategies: List[ComparisonPromptStrategy] = [
            SparkPlugComparisonStrategy(),
            BrakeDiscComparisonStrategy(),
            HornComparisonStrategy(),
        ]
        self._comparison_strategies: Dict[str, ComparisonPromptStrategy] = {
            strategy.part_key: strategy for strategy in strategies
        }
        self._comparison_alias_map: Dict[str, str] = {}
        for strategy in strategies:
            self._comparison_alias_map[strategy.part_key] = strategy.part_key
            for alias in strategy.aliases:
                self._comparison_alias_map[alias] = strategy.part_key

    @staticmethod
    def _extract_docx_text(file_bytes: bytes) -> str:
        """Extract meaningful text (including tables) from a DOCX file."""
        document = Document(io.BytesIO(file_bytes))
        lines: List[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)

        for table in document.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    lines.append(" | ".join(row_cells))

        combined = "\n".join(lines).strip()
        return combined if combined else "RFQ document (DOCX) contained no extractable text."

    def _prepare_rfq_input(self, rfq_bytes: bytes, rfq_mime: str) -> Tuple[Union[bytes, str], Optional[str]]:
        """Prepare RFQ input for Gemini, converting DOCX files to plain text."""
        if rfq_mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                text_content = self._extract_docx_text(rfq_bytes)
                logger.info("[COMPARE] Converted DOCX RFQ to plain text for Gemini input")
                return text_content, None
            except Exception as exc:
                logger.error("[COMPARE] Failed to read DOCX RFQ", exc_info=True)
                raise HTTPException(
                    status_code=400,
                    detail="Unable to read DOCX file. Please ensure the document is not corrupted."
                ) from exc
        if rfq_mime == "application/msword":
            raise HTTPException(
                status_code=400,
                detail="Legacy .doc files are not supported. Please upload the RFQ as PDF or convert it to DOCX."
            )
        return rfq_bytes, rfq_mime

    @staticmethod
    def _strip_markdown_fence(text: str) -> str:
        """Remove surrounding markdown fences (``` or ```json)."""
        text = text.strip()
        if text.startswith("```json"):
            text = text[7:]
        elif text.startswith("```"):
            text = text[3:]
        if text.endswith("```"):
            text = text[:-3]
        return text.strip()

    @staticmethod
    def _quote_unquoted_keys(text: str) -> str:
        """Ensure object keys are wrapped in double quotes for valid JSON parsing."""
        pattern = re.compile(r'(?m)^(?P<indent>\s*)(?P<key>[A-Za-z][A-Za-z0-9_\-\s]*?)\s*:(?=\s)')

        def replacer(match: re.Match) -> str:
            indent = match.group("indent")
            key = match.group("key").strip()
            # Skip keys that are already quoted
            if key.startswith('"') and key.endswith('"'):
                return match.group(0)
            return f'{indent}"{key}":'

        return pattern.sub(replacer, text)

    @staticmethod
    def _clean_list_item(value: str) -> str:
        cleaned = value.strip().strip(",")
        if cleaned.startswith('"') and cleaned.endswith('"'):
            cleaned = cleaned[1:-1]
        return cleaned.strip()

    def _extract_list_section(self, text: str, key: str) -> List[str]:
        """Extract an array section even if formatting is slightly invalid."""
        items: List[str] = []
        # Primary: look for JSON-style array
        array_pattern = re.compile(
            rf'\b{key}\b\s*[:=]\s*\[(?P<body>.*?)\]',
            re.IGNORECASE | re.DOTALL,
        )
        match = array_pattern.search(text)
        if match:
            body = match.group("body")
            # Prefer quoted strings
            quoted_items = re.findall(r'"([^"]+)"', body)
            if quoted_items:
                items.extend(item.strip() for item in quoted_items if item.strip())
            else:
                # Fall back to comma-separated entries
                for part in body.split(","):
                    cleaned = self._clean_list_item(part)
                    if cleaned:
                        items.append(cleaned)
            if items:
                return items

        # Secondary: capture bullet list following the key
        lines = text.replace("\r", "").split("\n")
        collecting = False
        for line in lines:
            stripped = line.strip()
            if not collecting:
                if re.match(rf'^\s*{key}\s*[:=]?\s*$', stripped, re.IGNORECASE):
                    collecting = True
                continue

            if not stripped:
                break

            if re.match(r'^[A-Za-z0-9_]+\s*[:=]', stripped):
                # Next top-level field detected
                break

            if stripped.startswith(("-", "*")):
                stripped = stripped[1:].strip()

            cleaned = self._clean_list_item(stripped)
            if cleaned:
                items.append(cleaned)

        return items

    def _parse_comparison_fallback(self, text: str):
        """Fallback parser when JSON decoding fails."""
        logger.info("[COMPARISON PARSER] Attempting fallback parsing heuristics")
        normalized = text.replace("\r", "")

        result = {
            "match": False,
            "confidence": "",
            "summary": "",
            "rfq_requirements": [],
            "cad_findings": [],
            "mismatches": [],
            "recommendations": "",
        }

        match_match = re.search(r'\bmatch\b\s*[:=]\s*(true|false)', normalized, re.IGNORECASE)
        if match_match:
            result["match"] = match_match.group(1).lower() == "true"

        confidence_match = re.search(r'\bconfidence\b\s*[:=]\s*"?(High|Medium|Low)"?', normalized, re.IGNORECASE)
        if confidence_match:
            result["confidence"] = confidence_match.group(1).capitalize()

        summary_match = re.search(r'\bsummary\b\s*[:=]\s*"([^"]+)"', normalized)
        if summary_match:
            result["summary"] = summary_match.group(1).strip()
        else:
            summary_match = re.search(r'\bsummary\b\s*[:=]\s*(.+)', normalized)
            if summary_match:
                result["summary"] = summary_match.group(1).strip().rstrip(",")

        for section in ("rfq_requirements", "cad_findings", "mismatches"):
            extracted = self._extract_list_section(normalized, section)
            if extracted:
                result[section] = extracted

        rec_list = self._extract_list_section(normalized, "recommendations")
        if rec_list:
            result["recommendations"] = "; ".join(rec_list)
        else:
            rec_match = re.search(r'\brecommendations?\b\s*[:=]\s*"([^"]+)"', normalized, re.IGNORECASE)
            if rec_match:
                result["recommendations"] = rec_match.group(1).strip()
            else:
                rec_match = re.search(r'\brecommendations?\b\s*[:=]\s*(.+)', normalized, re.IGNORECASE)
                if rec_match:
                    result["recommendations"] = rec_match.group(1).strip()

        if any([result["summary"], result["rfq_requirements"], result["cad_findings"], result["mismatches"], result["recommendations"]]):
            logger.info("[COMPARISON PARSER] Fallback parsing produced usable data")
            return result

        logger.error("[COMPARISON PARSER] Fallback parsing failed to extract data")
        return None

    def inspect_drawing(self, file_bytes: bytes, mime_type: str) -> str:
        """Analyze a CAD or welding drawing (image or PDF) and generate a detailed report."""
        logger.info(f"=== START INSPECTION === (MIME type: {mime_type})")

        prompt = (
            "You are an expert welding engineer analyzing a technical CAD or welding drawing. "
            "Carefully examine the entire drawing and identify ALL welds present in the diagram.\n\n"
            "INSTRUCTIONS:\n"
            "1. Identify every weld symbol, weld callout, and welding annotation visible in the drawing.\n"
            "2. Assign sequential serial numbers starting from W1, W2, W3, and continue for ALL welds found.\n"
            "3. Extract part numbers, plate numbers (PL), component numbers, and any other identifiers EXACTLY as shown in the drawing.\n"
            "4. Do NOT assume or invent part numbers - use only what is clearly visible in the drawing.\n"
            "5. Identify all welds regardless of their type or location.\n\n"
            "For each weld identified, provide:\n"
            "- Serial No (W1, W2, W3, ... - assign sequentially for all welds found)\n"
            "- Description (detailed description using EXACT part numbers, plate numbers, and component names as shown in the drawing)\n"
            "- Welding Type (e.g., Fillet Weld, Double Fillet Weld, Groove Weld, etc. - identify from weld symbols)\n"
            "- Welding Value (size in mm - extract exactly as specified in the drawing)\n"
            "- Remarks (any notes, flags, or annotations like TYP, OF DRIVE, OF MOTOR, M, etc. - use exactly as shown)\n"
            "- Position (location description referencing the exact parts, sections, or views shown)\n"
            "- Confidence (High if clearly visible, Medium if partially visible, Low if uncertain)\n\n"
            "CRITICAL OUTPUT REQUIREMENTS:\n"
            "1. Output ONLY valid JSON format with NO additional text before or after.\n"
            "2. The JSON structure must be:\n"
            "   {\n"
            "     \"welds\": [\n"
            "       {\n"
            "         \"Serial No\": \"W1\",\n"
            "         \"Description\": \"...\",\n"
            "         \"Welding Type\": \"...\",\n"
            "         \"Welding Value\": \"...\",\n"
            "         \"Remarks\": \"...\",\n"
            "         \"Position\": \"...\",\n"
            "         \"Confidence\": \"...\"\n"
            "       },\n"
            "       ...\n"
            "     ],\n"
            "     \"explanations\": \"Detailed explanations for each weld...\"\n"
            "   }\n"
            "3. Include ALL welds found - do not stop early. Continue until you have identified every weld in the drawing.\n"
            "4. Use EXACT part numbers, plate numbers, and identifiers as they appear in the drawing (e.g., Part 1, PL10-21, Beam B11, etc.).\n"
            "5. The \"welds\" array should contain one object per weld with all 7 fields.\n"
            "6. The \"explanations\" field should contain detailed explanations for each weld, referencing the exact locations and symbols from the drawing.\n"
            "7. Output ONLY valid JSON - no markdown, no code blocks, no explanations outside the JSON structure."
        )

        response_text = self.client.chat(file_bytes, mime_type, prompt)

        logger.info("=== INSPECTION COMPLETE ===")
        return response_text

    def _get_comparison_prompt(self, part: str) -> str:
        normalized = (part or "").strip().lower()
        strategy_key = self._comparison_alias_map.get(normalized)

        if not strategy_key:
            if normalized:
                logger.info("[COMPARE] Unknown part '%s'; defaulting to spark plug prompt", normalized)
            strategy_key = "spark_plug"

        strategy = self._comparison_strategies.get(strategy_key)
        if not strategy:
            logger.warning(
                "[COMPARE] Strategy '%s' not found; falling back to default spark plug prompt",
                strategy_key,
            )
            strategy = self._comparison_strategies["spark_plug"]

        logger.info("[COMPARE] Using comparison strategy: %s (for part: %s)", strategy_key, normalized)
        return strategy.build_prompt()

    def compare_rfq_and_cad(
        self,
        rfq_input: Tuple[Union[bytes, str], Optional[str]],
        cad_bytes: bytes,
        cad_mime: str,
        part: str = "spark_plug",
    ) -> str:
        """Compare RFQ document with CAD drawing and summarize alignment."""
        logger.info("=== START RFQ VS CAD COMPARISON === (part=%s)", part or "spark_plug")

        prompt = self._get_comparison_prompt(part)

        response_text = self.client.chat_with_files(
            prompt,
            [
                rfq_input,
                (cad_bytes, cad_mime),
            ],
        )

        logger.info("=== COMPARISON COMPLETE ===")
        return response_text

    @staticmethod
    def _normalize_value(value: str) -> str:
        """
        Normalize RFQ/CAD values so that formatting differences
        (spaces, commas, diameter symbol, units, etc.) don't cause false mismatches.

        Examples:
        "Ø300 mm"    -> "300"
        "Ø300"       -> "300"
        "75,5 mm"    -> "75.5"
        "4 x Ø23.5"  -> "4x23.5"
        "22.2°"      -> "22.2"
        "M14"        -> "m14"
        """
        if not value:
            return ""

        v = value.strip().lower()

        # 1) normalize decimal separator
        v = v.replace(",", ".")

        # 2) remove diameter & degree symbols
        v = v.replace("ø", "").replace("°", "")

        # 3) normalize multiplication/group notation: "4 x 23.5", "4-23.5" → "4x23.5"
        v = v.replace("×", "x")
        v = re.sub(r"\s*[-x]\s*", "x", v)

        # 4) strip common units
        units = [" mm", "mm", " in", "in"]
        for unit in units:
            if v.endswith(unit):
                v = v[: -len(unit)].strip()
                break

        # 5) if pure number -> canonical numeric form
        num_match = re.fullmatch(r"[-+]?\d*\.?\d+", v)
        if num_match:
            return num_match.group(0)

        # 6) fallback: compact spaces/commas on the cleaned string
        v = re.sub(r"[\s,]+", "", v)
        return v.strip()

    def _normalize_label(self, label: str) -> str:
        """
        Normalize parameter labels to create consistent dictionary keys for matching.
        
        Examples:
        "Thread Size"    -> "threadsize"
        "Thread_Size"    -> "threadsize"
        "Part Number"    -> "partnumber"
        "Hex Size"       -> "hexsize"
        """
        if not label:
            return ""
        
        # Convert to lowercase and strip whitespace
        normalized = label.strip().lower()
        
        # Replace common separators (spaces, underscores, hyphens) with nothing
        normalized = re.sub(r"[\s_\-]+", "", normalized)
        
        return normalized

    def _spec_list_to_dict(self, items: List[str]) -> Dict[str, Dict[str, str]]:
        """Convert ['Thread Size: M14', ...] to key→{label,value} dict."""
        specs: Dict[str, Dict[str, str]] = {}
        pattern = re.compile(r"^\s*([^:]+?)\s*[:\-]\s*(.+?)\s*$")
        for item in items:
            match = pattern.match(item)
            if not match:
                continue
            raw_label, raw_value = match.groups()
            key = self._normalize_label(raw_label)
            if not key:
                continue
            specs[key] = {"label": raw_label.strip(), "value": raw_value.strip()}
        return specs

    def build_metric_records(
        self,
        rfq_requirements: List[str],
        cad_findings: List[str],
    ) -> List[Dict]:
        """
        Build canonical metric records with backend-determined status.
        
        Returns a list of records with:
        {
            "key": normalized key (e.g., "threadsize"),
            "label": original RFQ label (e.g., "Thread Size"),
            "rfq_value": RFQ value string,
            "cad_value": CAD value string (or "" if missing),
            "match": "Match" | "Mismatch" | "Missing" | "Extra"
        }
        """
        rfq_specs = self._spec_list_to_dict(rfq_requirements)
        cad_specs = self._spec_list_to_dict(cad_findings)
        
        if not rfq_specs:
            logger.info("[METRIC-RECORDS] No RFQ specs available")
            return []
        
        records: List[Dict] = []
        
        # Process RFQ metrics
        for key, rfq_spec in rfq_specs.items():
            cad_spec = cad_specs.get(key)
            cad_value = cad_spec.get("value", "").strip() if cad_spec else ""
            rfq_value = rfq_spec.get("value", "").strip()
            
            # Determine status
            if not cad_spec or not cad_value:
                match_status = "Missing"
            elif self._normalize_value(rfq_value) == self._normalize_value(cad_value):
                match_status = "Match"
            else:
                match_status = "Mismatch"
            
            records.append({
                "key": key,
                "label": rfq_spec.get("label", ""),
                "rfq_value": rfq_value,
                "cad_value": cad_value,
                "match": match_status,
            })
        
        # Optionally include CAD-only metrics as "Extra"
        for key, cad_spec in cad_specs.items():
            if key not in rfq_specs:
                records.append({
                    "key": key,
                    "label": cad_spec.get("label", ""),
                    "rfq_value": "",
                    "cad_value": cad_spec.get("value", "").strip(),
                    "match": "Extra",
                })
        
        logger.info("[METRIC-RECORDS] Built %d metric records", len(records))
        return records

    def _extract_cad_bboxes(
        self,
        cad_bytes: bytes,
        cad_mime: str,
        metric_records: List[Dict],
    ) -> List[Dict]:
        """
        Ask Gemini for bounding boxes ONLY (no status/value decisions).
        
        Returns a list of entries such as:
        {
          "parameter": "Thread Length",
          "key": "threadlength",
          "bounding_box": [x1, y1, x2, y2]
        }
        """
        # Build text block of metrics that have CAD values
        metrics_with_cad = [
            f"{record['label']}: {record['cad_value']}"
            for record in metric_records
            if record.get("cad_value") and record.get("match") != "Missing"
        ]
        
        if not metrics_with_cad:
            logger.info("[ANNOTATION] No metrics with CAD values to locate")
            return []
        
        metrics_text = "\n".join(metrics_with_cad)
        
        prompt = (
            "You are an expert in reading engineering drawings.\n"
            "You will receive:\n"
            "1) A list of metric names and their values (already determined from the drawing).\n"
            "2) A CAD drawing image.\n\n"
            "METRICS TO LOCATE (name: value):\n"
            f"{metrics_text}\n\n"
            "YOUR TASK:\n"
            "For each metric in the list above, locate the corresponding dimension text/annotation on the drawing.\n"
            "Draw a bounding box around the dimension text and any associated arrows or callouts.\n\n"
            "CRITICAL RULES:\n"
            "  - Use the metric name EXACTLY as provided in the \"parameter\" field.\n"
            "  - Do NOT change or re-interpret the metric values.\n"
            "  - Do NOT decide if values match or mismatch - that is already determined.\n"
            "  - Only provide bounding boxes for dimensions you can clearly see on the drawing.\n"
            "  - If you cannot find a metric, simply omit it from the output.\n"
            "  - The bounding_box MUST be [x_min, y_min, x_max, y_max] with all values between 0 and 1,\n"
            "    normalized relative to the image width/height (0 = left/top, 1 = right/bottom).\n\n"
            "OUTPUT FORMAT:\n"
            "Return ONLY a JSON array, no extra text, no markdown.\n"
            "Example:\n"
            "[\n"
            "  {\n"
            '    "parameter": "Thread Length",\n'
            '    "bounding_box": [0.1, 0.2, 0.16, 0.22]\n'
            "  },\n"
            "  {\n"
            '    "parameter": "Overall Length",\n'
            '    "bounding_box": [0.25, 0.08, 0.32, 0.11]\n'
            "  }\n"
            "]\n"
        )

        try:
            response_text = self.client.chat_with_files(
                prompt,
                [
                    (cad_bytes, cad_mime),
                ],
            )

            cleaned = self._strip_markdown_fence(response_text).strip()

            if not cleaned:
                logger.warning("[ANNOTATION] Gemini returned empty bbox payload")
                return []

            parsed: Optional[Union[List, Dict]] = None

            try:
                parsed = json.loads(cleaned)
            except json.JSONDecodeError as e:
                logger.debug("[ANNOTATION] Initial JSON parse failed (may be truncated), trying fallback strategies")
                
                # Strategy 1: Try to extract complete JSON objects from a potentially truncated array
                # Find all complete JSON objects in the text
                objects = []
                depth = 0
                start = -1
                in_string = False
                escape = False
                
                for i, char in enumerate(cleaned):
                    if escape:
                        escape = False
                        continue
                    if char == '\\':
                        escape = True
                        continue
                    if char == '"' and not escape:
                        in_string = not in_string
                        continue
                    if in_string:
                        continue
                    
                    if char == '{':
                        if depth == 0:
                            start = i
                        depth += 1
                    elif char == '}':
                        depth -= 1
                        if depth == 0 and start >= 0:
                            # Found a complete object
                            try:
                                obj_str = cleaned[start:i+1]
                                obj = json.loads(obj_str)
                                if isinstance(obj, dict):
                                    objects.append(obj)
                            except json.JSONDecodeError:
                                pass
                            start = -1
                
                if objects:
                    logger.info(
                        "[ANNOTATION] Extracted %d complete objects from truncated JSON",
                        len(objects)
                    )
                    parsed = objects
                else:
                    # Strategy 2: Try to find JSON array or object with regex (less reliable)
                    array_match = re.search(r'\[.*\]', cleaned, re.DOTALL)
                    object_match = re.search(r'\{.*\}', cleaned, re.DOTALL)
                    candidate = (
                        array_match.group(0)
                        if array_match
                        else object_match.group(0)
                        if object_match
                        else None
                    )

                    if candidate:
                        try:
                            parsed = json.loads(candidate)
                        except json.JSONDecodeError:
                            logger.warning(
                                "[ANNOTATION] Failed to parse candidate JSON chunk (response likely truncated at MAX_TOKENS)"
                            )
                    else:
                        logger.warning("[ANNOTATION] No JSON structure detected in bbox response")

            if parsed is None:
                logger.warning("[ANNOTATION] Gemini bbox response was not valid JSON (response may be truncated)")
                return []

            # Support both raw list and {"annotations": [...]} style
            if isinstance(parsed, dict) and "annotations" in parsed:
                data = parsed.get("annotations")
            else:
                data = parsed

            if not isinstance(data, list):
                logger.warning("[ANNOTATION] Bbox payload is not an array")
                return []

            bbox_entries: List[Dict] = []
            for entry in data:
                if not isinstance(entry, dict):
                    continue
                parameter = str(entry.get("parameter", "")).strip()
                bbox = entry.get("bounding_box") or entry.get("bbox")
                if not parameter:
                    continue
                
                # Validate bounding box format
                if bbox is not None:
                    if not isinstance(bbox, list) or len(bbox) != 4:
                        logger.debug(
                            "[ANNOTATION] Skipping invalid bbox format for parameter '%s': %s",
                            parameter,
                            bbox
                        )
                        continue
                    # Ensure all values are numeric
                    try:
                        bbox = [float(x) for x in bbox]
                    except (ValueError, TypeError):
                        logger.debug(
                            "[ANNOTATION] Skipping bbox with non-numeric values for parameter '%s': %s",
                            parameter,
                            bbox
                        )
                        continue
                
                # Map parameter back to metric record to get normalized key
                normalized_key = self._normalize_label(parameter)
                
                bbox_entries.append({
                    "parameter": parameter,
                    "key": normalized_key,
                    "bounding_box": bbox,
                })

            logger.info("[ANNOTATION] Extracted %d CAD bounding boxes", len(bbox_entries))
            return bbox_entries
            
        except Exception as exc:
            logger.warning("[ANNOTATION] Unable to extract CAD bounding boxes: %s", exc, exc_info=True)
            return []

    @staticmethod
    def _normalize_bbox(
        bbox: List[float],
        width: int,
        height: int,
    ) -> Optional[Tuple[int, int, int, int]]:
        """
        Normalize model bbox into (x1, y1, x2, y2) in image pixels.
        Returns None for invalid boxes.
        """
        if len(bbox) != 4:
            return None

        x1, y1, x2, y2 = [float(b) for b in bbox]
        img_max = float(max(width, height))

        # --- Case A: normalized [0..1] ---
        if 0 <= x1 <= 1 and 0 <= y1 <= 1 and 0 <= x2 <= 1 and 0 <= y2 <= 1:
            x1 *= width
            x2 *= width
            y1 *= height
            y2 *= height

        # --- Case B: percent [0..100] ---
        elif 0 <= x1 <= 100 and 0 <= y1 <= 100 and 0 <= x2 <= 100 and 0 <= y2 <= 100:
            x1 = (x1 / 100) * width
            x2 = (x2 / 100) * width
            y1 = (y1 / 100) * height
            y2 = (y2 / 100) * height

        # --- Reject boxes excessively outside image (hallucinated) ---
        if max(abs(x1), abs(y1), abs(x2), abs(y2)) > img_max * 2:
            return None

        # --- Handle width/height format [x, y, w, h] ---
        if x2 > 0 and y2 > 0 and (x2 <= width and y2 <= height) and (x2 < x1 or y2 < y1):
            w, h = x2, y2
            x2 = x1 + w
            y2 = y1 + h

        # --- Clamp ---
        x1 = max(0, min(int(round(x1)), width - 1))
        x2 = max(0, min(int(round(x2)), width - 1))
        y1 = max(0, min(int(round(y1)), height - 1))
        y2 = max(0, min(int(round(y2)), height - 1))

        # --- Reject inverted / zero-area boxes ---
        if x2 <= x1 or y2 <= y1:
            return None

        # --- Reject very tiny boxes ---
        if (x2 - x1) < 3 or (y2 - y1) < 3:
            return None

        # --- Reject extremely large boxes (almost whole image) ---
        box_area = (x2 - x1) * (y2 - y1)
        img_area = width * height
        if box_area > img_area * 0.80:
            return None

        return x1, y1, x2, y2

    def _clamp_bbox(bbox: List[float], width: int, height: int) -> Tuple[int, int, int, int]:
        x1, y1, x2, y2 = bbox
        x1 = max(0, min(int(round(x1)), width - 1))
        x2 = max(0, min(int(round(x2)), width - 1))
        y1 = max(0, min(int(round(y1)), height - 1))
        y2 = max(0, min(int(round(y2)), height - 1))
        if x2 < x1:
            x1, x2 = x2, x1
        if y2 < y1:
            y1, y2 = y2, y1
        return x1, y1, x2, y2

    def _annotate_cad_image(
        self,
        cad_bytes: bytes,
        comparisons: List[Dict],
    ) -> Optional[str]:
        """Create annotated CAD image highlighting matches/mismatches."""
        np_bytes = np.frombuffer(cad_bytes, np.uint8)
        image = cv2.imdecode(np_bytes, cv2.IMREAD_COLOR)
        if image is None:
            logger.warning("[ANNOTATION] Unable to decode CAD image for annotation")
            return None

        height, width = image.shape[:2]
        colors = {
            "Match": (0, 180, 0),
            "Mismatch": (0, 0, 255),
            "Missing": (0, 215, 255),
            "Extra": (255, 165, 0),  # Orange for CAD-only metrics
        }

        for record in comparisons:
            bbox = record.get("bounding_box")
            if not bbox or len(bbox) != 4:
                continue
            norm = self._normalize_bbox(bbox, width, height)
            if not norm:
                logger.debug(f"[ANNOTATION] Skipping invalid bbox: {bbox}")
                continue

            x1, y1, x2, y2 = norm
            color = colors.get(record.get("match", ""), (255, 255, 255))
            cv2.rectangle(image, (x1, y1), (x2, y2), color, 2)
            label_text = f"{record.get('parameter', '')}: {record.get('match', '')}"
            cv2.putText(
                image,
                label_text,
                (x1, max(y1 - 8, 16)),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.5,
                color,
                2,
                cv2.LINE_AA,
            )

        success, buffer = cv2.imencode(".png", image)
        if not success:
            logger.warning("[ANNOTATION] Failed to encode annotated image")
            return None
        encoded = base64.b64encode(buffer).decode("utf-8")
        return f"data:image/png;base64,{encoded}"

    def generate_auto_annotations(
        self,
        rfq_requirements: List[str],
        cad_findings: List[str],
        cad_bytes: bytes,
        cad_mime: str,
    ) -> Tuple[Optional[str], List[Dict]]:
        """
        Build comparison records and annotated image using two-step flow:
        1. build_metric_records: backend determines Match/Mismatch/Missing status
        2. _extract_cad_bboxes: Gemini only provides bounding box locations
        """
        if not cad_mime.startswith("image/"):
            logger.info("[ANNOTATION] CAD file is not an image; skipping auto-annotation")
            return None, []

        # Step 1: Build canonical metric records with backend-determined status
        try:
            metric_records = self.build_metric_records(rfq_requirements, cad_findings)
        except Exception as exc:
            logger.warning("[ANNOTATION] Unable to build metric records: %s", exc, exc_info=True)
            return None, []

        if not metric_records:
            logger.info("[ANNOTATION] No metric records available for annotation")
            return None, []

        # Step 2: Get bounding boxes from Gemini (no status/value decisions)
        try:
            bbox_entries = self._extract_cad_bboxes(cad_bytes, cad_mime, metric_records)
        except Exception as exc:
            logger.warning("[ANNOTATION] Unable to extract CAD bounding boxes: %s", exc, exc_info=True)
            bbox_entries = []

        # Step 3: Merge metric records with bounding boxes
        bbox_lookup = {entry.get("key"): entry for entry in bbox_entries}
        
        comparison_records: List[Dict] = []
        for record in metric_records:
            bbox_entry = bbox_lookup.get(record.get("key"))
            bbox = bbox_entry.get("bounding_box") if bbox_entry else None
            
            # Normalize bounding box format
            if isinstance(bbox, list) and len(bbox) == 4:
                bbox_values = [float(coord) for coord in bbox]
            else:
                bbox_values = None

            comparison_records.append({
                "parameter": record.get("label", ""),
                "rfq_value": record.get("rfq_value", ""),
                "cad_value": record.get("cad_value", ""),
                "match": record.get("match", "Missing"),
                "bounding_box": bbox_values,
            })

        # Step 4: Create annotated image with only records that have bounding boxes
        annotated_image = self._annotate_cad_image(
            cad_bytes,
            [record for record in comparison_records if record.get("bounding_box")],
        )

        logger.info(
            "[ANNOTATION] Generated %d annotation records (%d with bboxes)",
            len(comparison_records),
            len([r for r in comparison_records if r.get("bounding_box")]),
        )

        return annotated_image, comparison_records

    def parse_json_response(self, response_text: str):
        """Parse JSON response from Gemini and extract welds and explanations.
        
        Returns:
            tuple: (welds_list, explanations_string) or (None, None) if parsing fails
        """
        original_response = response_text
        try:
            logger.info(f"[JSON PARSER] Starting JSON parsing. Raw response length: {len(response_text)} chars")
            logger.info(f"[JSON PARSER] Raw response (first 500 chars): {response_text[:500]}")
            
            # Try to extract JSON from response (might have code blocks or extra text)
            response_text = response_text.strip()
            
            # Remove markdown code blocks if present
            had_markdown = False
            if response_text.startswith("```json"):
                response_text = response_text[7:]  # Remove ```json
                had_markdown = True
                logger.info("[JSON PARSER] Removed ```json markdown block")
            elif response_text.startswith("```"):
                response_text = response_text[3:]  # Remove ```
                had_markdown = True
                logger.info("[JSON PARSER] Removed ``` markdown block")
            
            if response_text.endswith("```"):
                response_text = response_text[:-3]  # Remove closing ```
                had_markdown = True
                logger.info("[JSON PARSER] Removed closing ``` markdown block")
            
            response_text = response_text.strip()
            
            if had_markdown:
                logger.info(f"[JSON PARSER] After markdown removal (first 500 chars): {response_text[:500]}")
            
            # Find JSON object in the text
            json_start = response_text.find("{")
            json_end = response_text.rfind("}") + 1
            
            if json_start == -1 or json_end == 0:
                logger.warning("⚠ No JSON object found in response.")
                logger.warning(f"[JSON PARSER] Response text (first 1000 chars): {original_response[:1000]}")
                return None, None
            
            json_str = response_text[json_start:json_end]
            logger.info(f"[JSON PARSER] Extracted JSON string length: {len(json_str)} chars")
            logger.info(f"[JSON PARSER] Extracted JSON (first 500 chars): {json_str[:500]}")
            
            # Fix common JSON syntax errors from LLM before parsing
            # Fix 1: Remove period before colon (e.g., "key":. "value" -> "key": "value")
            original_json = json_str
            json_str = re.sub(r':\s*\.\s*"', ': "', json_str)
            json_str = re.sub(r':\s*\.\s*(\d+)', r': \1', json_str)
            json_str = re.sub(r':\s*\.\s*([A-Z])', r': "\1', json_str)
            if json_str != original_json:
                logger.info("[JSON PARSER] Fixed period-before-colon syntax errors")
            
            # Try to parse JSON with better error handling
            try:
                data = json.loads(json_str)
            except json.JSONDecodeError as parse_error:
                # Log the error position and surrounding context
                error_pos = getattr(parse_error, 'pos', None)
                if error_pos:
                    logger.error(f"[JSON PARSER] JSON decode error at position {error_pos}")
                    # Show context around the error
                    start = max(0, error_pos - 100)
                    end = min(len(json_str), error_pos + 100)
                    logger.error(f"[JSON PARSER] Context around error: ...{json_str[start:end]}...")
                    logger.error(f"[JSON PARSER] Error line/column: line {getattr(parse_error, 'lineno', '?')}, col {getattr(parse_error, 'colno', '?')}")
                
                # Try to extract welds array even if full JSON is broken
                logger.warning("[JSON PARSER] Attempting to extract welds array from broken JSON...")
                
                # Find the welds array and extract individual weld objects
                welds_start = json_str.find('"welds": [')
                if welds_start >= 0:
                    # Extract welds array content
                    array_start = json_str.find('[', welds_start)
                    bracket_count = 1
                    array_end = -1
                    for i in range(array_start + 1, min(len(json_str), array_start + 50000)):  # Limit search
                        if json_str[i] == '[':
                            bracket_count += 1
                        elif json_str[i] == ']':
                            bracket_count -= 1
                            if bracket_count == 0:
                                array_end = i + 1
                                break
                    
                    if array_end > array_start:
                        welds_array_str = json_str[array_start:array_end]
                        try:
                            welds_array = json.loads(welds_array_str)
                            data = {"welds": welds_array, "explanations": ""}
                            logger.info(f"[JSON PARSER] ✅ Successfully extracted {len(welds_array)} welds from broken JSON")
                        except:
                            # Try extracting individual weld objects using regex
                            weld_objects = []
                            # Pattern to match a complete weld object (from Serial No to Confidence)
                            weld_pattern = r'\{\s*"Serial No"[^}]*"Confidence"[^}]*\}'
                            matches = re.finditer(weld_pattern, welds_array_str, re.DOTALL)
                            for match in matches:
                                try:
                                    weld_obj = json.loads(match.group(0))
                                    weld_objects.append(weld_obj)
                                except:
                                    pass
                            
                            if weld_objects:
                                data = {"welds": weld_objects, "explanations": ""}
                                logger.info(f"[JSON PARSER] ✅ Extracted {len(weld_objects)} welds using regex fallback")
                            else:
                                logger.error("[JSON PARSER] Could not extract any welds")
                                raise parse_error
                    else:
                        logger.error("[JSON PARSER] Could not find welds array end")
                        raise parse_error
                else:
                    logger.error("[JSON PARSER] Could not find welds array")
                    raise parse_error
            logger.info(f"[JSON PARSER] Successfully parsed JSON. Keys: {list(data.keys())}")
            
            # Extract welds array
            welds = data.get("welds", [])
            explanations = data.get("explanations", "")
            
            logger.info(f"[JSON PARSER] Found {len(welds)} welds in JSON response")
            
            if not welds:
                logger.warning("⚠ No welds found in JSON response.")
                logger.warning(f"[JSON PARSER] Full JSON data: {data}")
                return None, None
            
            # Convert to list of dictionaries with consistent keys
            result = []
            for idx, weld in enumerate(welds):
                # Ensure all expected keys exist
                weld_dict = {
                    "Serial No": str(weld.get("Serial No", "")),
                    "Description": str(weld.get("Description", "")),
                    "Welding Type": str(weld.get("Welding Type", "")),
                    "Welding Value": str(weld.get("Welding Value", "")),
                    "Remarks": str(weld.get("Remarks", "")),
                    "Position": str(weld.get("Position", "")),
                    "Confidence": str(weld.get("Confidence", ""))
                }
                result.append(weld_dict)
                if idx < 2:  # Log first 2 welds as sample
                    logger.info(f"[JSON PARSER] Sample weld {idx + 1}: {list(weld_dict.keys())}")
            
            logger.info(f"✅ Successfully parsed {len(result)} welds from JSON")
            logger.info(f"[JSON PARSER] First weld keys: {list(result[0].keys()) if result else 'N/A'}")
            return result, explanations
            
        except json.JSONDecodeError as e:
            logger.error(f"⚠ JSON parsing error: {e}")
            logger.error(f"[JSON PARSER] Error position: {e.pos if hasattr(e, 'pos') else 'N/A'}")
            logger.error(f"[JSON PARSER] Response text (first 1000 chars): {original_response[:1000]}")
            logger.error(f"[JSON PARSER] Full response length: {len(original_response)} chars")
            return None, None

    def parse_vendor_comparison_response(self, response_text: str):
        """Parse vendor comparison JSON response from Gemini."""
        original_response = response_text
        try:
            response_text = self._strip_markdown_fence(response_text)
            normalized_text = self._quote_unquoted_keys(response_text)

            json_start = normalized_text.find("{")
            json_end = normalized_text.rfind("}") + 1

            if json_start == -1 or json_end <= json_start:
                logger.error("[VENDOR-COMPARISON PARSER] JSON object not found")
                return None

            json_str = normalized_text[json_start:json_end]

            data = json.loads(json_str)
            
            logger.info("[VENDOR-COMPARISON PARSER] Parsed JSON keys: %s", list(data.keys()))
            
            # Validate structure
            vendors = data.get("vendors", [])
            comparison = data.get("comparison", {})
            
            if not isinstance(vendors, list):
                logger.error("[VENDOR-COMPARISON PARSER] 'vendors' is not a list")
                return None
            
            result = {
                "vendors": vendors,
                "comparison": comparison,
            }

            logger.info(
                "[VENDOR-COMPARISON PARSER] Extracted: vendors=%d",
                len(vendors),
            )

            logger.info("[VENDOR-COMPARISON PARSER] Parsed vendor comparison response successfully")
            return result

        except json.JSONDecodeError as exc:
            logger.error("[VENDOR-COMPARISON PARSER] JSON decode error", exc_info=True)
            logger.error(f"[VENDOR-COMPARISON PARSER] Response text (first 1000 chars): {original_response[:1000]}")
            return None
        except Exception:
            logger.error("[VENDOR-COMPARISON PARSER] Unexpected error", exc_info=True)
            logger.error(f"[VENDOR-COMPARISON PARSER] Response text (first 1000 chars): {original_response[:1000]}")
            return None

    def parse_comparison_response(self, response_text: str):
        """Parse comparison JSON response from Gemini."""
        original_response = response_text
        try:
            response_text = self._strip_markdown_fence(response_text)
            normalized_text = self._quote_unquoted_keys(response_text)

            json_start = normalized_text.find("{")
            json_end = normalized_text.rfind("}") + 1

            if json_start == -1 or json_end <= json_start:
                logger.error("[COMPARISON PARSER] JSON object not found")
                return None

            json_str = normalized_text[json_start:json_end]

            data = json.loads(json_str)
            
            logger.info("[COMPARISON PARSER] Parsed JSON keys: %s", list(data.keys()))
            
            # Handle old format with "metrics" array (convert to new format)
            if "metrics" in data and "rfq_requirements" not in data:
                logger.warning("[COMPARISON PARSER] Detected old format with 'metrics' array, converting to new format")
                metrics = data.get("metrics", [])
                rfq_requirements = []
                cad_findings = []
                mismatches_list = []
                
                for metric in metrics:
                    if isinstance(metric, dict):
                        metric_name = metric.get("metric", "")
                        rfq_val = metric.get("rfq", "")
                        cad_val = metric.get("cad", "")
                        status = metric.get("status", "")
                        
                        if metric_name:
                            rfq_requirements.append(f"{metric_name}: {rfq_val}" if rfq_val else f"{metric_name}: —")
                            cad_findings.append(f"{metric_name}: {cad_val}" if cad_val else f"{metric_name}: —")
                            
                            if status in ["Mismatch", "Missing in CAD"]:
                                mismatches_list.append(f"{metric_name}: RFQ {rfq_val} vs CAD {cad_val} - {status}")
                
                # Also check mismatches array if present
                if "mismatches" in data:
                    for mismatch in data.get("mismatches", []):
                        if isinstance(mismatch, dict):
                            metric_name = mismatch.get("metric", "")
                            rfq_val = mismatch.get("rfq", "")
                            cad_val = mismatch.get("cad", "")
                            status = mismatch.get("status", "")
                            if metric_name:
                                mismatches_list.append(f"{metric_name}: RFQ {rfq_val} vs CAD {cad_val} - {status}")
                        elif isinstance(mismatch, str):
                            mismatches_list.append(mismatch)
                
                result = {
                    "match": bool(data.get("match", False)),
                    "confidence": str(data.get("confidence", "")),
                    "summary": str(data.get("summary", "")),
                    "rfq_requirements": rfq_requirements,
                    "cad_findings": cad_findings,
                    "mismatches": mismatches_list,
                    "recommendations": str(data.get("recommendations", "")),
                }
            else:
                # Standard format
                result = {
                    "match": bool(data.get("match", False)),
                    "confidence": str(data.get("confidence", "")),
                    "summary": str(data.get("summary", "")),
                    "rfq_requirements": [str(item) for item in data.get("rfq_requirements", [])],
                    "cad_findings": [str(item) for item in data.get("cad_findings", [])],
                    "mismatches": [str(item) for item in data.get("mismatches", [])],
                    "recommendations": str(data.get("recommendations", "")),
                }
            
            logger.info(
                "[COMPARISON PARSER] Extracted: rfq_count=%d, cad_count=%d, mismatch_count=%d",
                len(result["rfq_requirements"]),
                len(result["cad_findings"]),
                len(result["mismatches"]),
            )
            
            if len(result["rfq_requirements"]) == 0 and len(result["cad_findings"]) == 0:
                logger.warning(
                    "[COMPARISON PARSER] WARNING: Both rfq_requirements and cad_findings are empty! "
                    "This may indicate the response format doesn't match expectations."
                )

            logger.info("[COMPARISON PARSER] Parsed comparison response successfully")
            return result

        except json.JSONDecodeError as exc:
            logger.error("[COMPARISON PARSER] JSON decode error", exc_info=True)
            logger.error(f"[COMPARISON PARSER] Response text (first 1000 chars): {original_response[:1000]}")
            fallback_result = self._parse_comparison_fallback(response_text)
            if fallback_result:
                return fallback_result
            return None
        except Exception:
            logger.error("[COMPARISON PARSER] Unexpected error", exc_info=True)
            logger.error(f"[COMPARISON PARSER] Response text (first 1000 chars): {original_response[:1000]}")
            fallback_result = self._parse_comparison_fallback(response_text)
            if fallback_result:
                return fallback_result
            return None


# Initialize client and inspector
client = GeminiClient(PROJECT, REGION, MODEL)
inspector = WeldingInspector(client)


@app.get("/")
def root():
    return {"message": "Welding Inspector API", "status": "running"}

@app.get("/health")
def health():
    """Health check endpoint."""
    return {"status": "healthy", "service": "Welding Inspector API"}


@app.post("/analyze")
async def analyze_image(file: UploadFile = File(...)):
    """Upload an image or PDF and analyze it for welding information."""
    try:
        # Validate file type - accept images and PDFs
        if not file.content_type:
            raise HTTPException(status_code=400, detail="File type not specified")
        
        is_image = file.content_type.startswith("image/")
        is_pdf = file.content_type == "application/pdf"
        
        if not (is_image or is_pdf):
            raise HTTPException(
                status_code=400, 
                detail=f"File must be an image (PNG, JPG, JPEG, WEBP) or PDF. Received: {file.content_type}"
            )

        # Read file bytes
        file_bytes = await file.read()
        
        # Validate file size (max 20MB)
        max_size = 20 * 1024 * 1024  # 20MB
        if len(file_bytes) > max_size:
            raise HTTPException(
                status_code=400,
                detail=f"File size exceeds maximum allowed size of 20MB. Received: {len(file_bytes) / (1024 * 1024):.2f}MB"
            )
        
        if len(file_bytes) == 0:
            raise HTTPException(status_code=400, detail="File is empty")
        
        logger.info(f"Processing file: {file.filename}, type: {file.content_type}, size: {len(file_bytes)} bytes")
        
        # Detect MIME type
        mime_type = file.content_type or mimetypes.guess_type(file.filename or "")[0] or "image/jpeg"
        
        # For PDFs, ensure correct MIME type
        if is_pdf and mime_type != "application/pdf":
            mime_type = "application/pdf"

        # Analyze the file (image or PDF)
        report = inspector.inspect_drawing(file_bytes, mime_type)
        
        logger.info(f"[ENDPOINT] Raw LLM report length: {len(report)} chars")
        logger.info(f"[ENDPOINT] Raw LLM report (first 500 chars): {report[:500]}")
        
        # Parse JSON response
        table_data, explanations = inspector.parse_json_response(report)
        
        if table_data is None:
            # Fallback: try to parse as markdown table if JSON parsing fails
            logger.warning("[ENDPOINT] JSON parsing failed, attempting markdown fallback")
            table_data = inspector.parse_table(report) if hasattr(inspector, 'parse_table') else None
            if "EXPLANATIONS:" in report:
                explanations = report.split("EXPLANATIONS:", 1)[1].strip()
        
        if table_data is not None:
            logger.info(f"[ENDPOINT] Parsed table_data: {len(table_data)} rows")
            if len(table_data) > 0:
                logger.info(f"[ENDPOINT] Table keys: {list(table_data[0].keys())}")
                logger.info(f"[ENDPOINT] Sample row: {table_data[0] if table_data else 'N/A'}")
        else:
            logger.warning("[ENDPOINT] ⚠ table_data is None - no table data to send to frontend")
        
        response_data = {
            "success": True,
            "report": report,
            "table": table_data,
            "explanations": explanations or ""
        }
        
        logger.info(f"[ENDPOINT] Sending response: success={response_data['success']}, table_rows={len(table_data) if table_data else 0}, has_explanations={bool(explanations)}")
        
        return JSONResponse(response_data)

    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        logger.error(f"Error analyzing file: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail=f"Error analyzing file: {str(e)}"
        )


@app.post("/compare")
async def compare_rfq_cad(
    rfq: UploadFile = File(...),
    cad: UploadFile = File(...),
    part: str = Form("spark_plug"),
):
    """Upload RFQ (PDF) and CAD (image/PDF) to compare alignment."""
    try:
        allowed_rfq_types = {
            "application/pdf",
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }

        if not rfq.content_type or rfq.content_type not in allowed_rfq_types:
            raise HTTPException(
                status_code=400,
                detail="RFQ file must be a PDF or Word document (.doc, .docx)",
            )

        if not cad.content_type:
            raise HTTPException(status_code=400, detail="CAD file type not specified")

        is_cad_image = cad.content_type.startswith("image/")
        is_cad_pdf = cad.content_type == "application/pdf"

        if not (is_cad_image or is_cad_pdf):
            raise HTTPException(
                status_code=400,
                detail="CAD file must be an image (PNG, JPG, JPEG, WEBP) or PDF",
            )

        rfq_bytes = await rfq.read()
        cad_bytes = await cad.read()

        if len(rfq_bytes) == 0 or len(cad_bytes) == 0:
            raise HTTPException(status_code=400, detail="Uploaded files cannot be empty")

        max_size = 20 * 1024 * 1024
        if len(rfq_bytes) > max_size or len(cad_bytes) > max_size:
            raise HTTPException(status_code=400, detail="Each file must be <= 20MB")

        rfq_mime = rfq.content_type if rfq.content_type in allowed_rfq_types else "application/pdf"
        cad_mime = cad.content_type if is_cad_image else "application/pdf"

        logger.info(
            "[COMPARE] Processing RFQ %s (%d bytes) and CAD %s (%d bytes)",
            rfq.filename,
            len(rfq_bytes),
            cad.filename,
            len(cad_bytes),
        )

        part_selection = (part or "spark_plug").strip() or "spark_plug"
        logger.info("[COMPARE] Part selection: %s", part_selection)

        rfq_input = inspector._prepare_rfq_input(rfq_bytes, rfq_mime)
        comparison_text = inspector.compare_rfq_and_cad(
            rfq_input,
            cad_bytes,
            cad_mime,
            part=part_selection,
        )

        result = inspector.parse_comparison_response(comparison_text)

        if result is None:
            logger.error("[COMPARE] Failed to parse comparison response for part: %s", part_selection)
            logger.error("[COMPARE] Raw response (first 2000 chars): %s", comparison_text[:2000] if comparison_text else "None")
            raise HTTPException(status_code=500, detail="Unable to parse comparison response")
        
        logger.info(
            "[COMPARE] Parsed result for part '%s': match=%s, rfq_count=%d, cad_count=%d, mismatch_count=%d",
            part_selection,
            result.get("match"),
            len(result.get("rfq_requirements", [])),
            len(result.get("cad_findings", [])),
            len(result.get("mismatches", [])),
        )

        annotated_image = None
        annotation_records: List[Dict] = []
        try:
            annotated_image, annotation_records = inspector.generate_auto_annotations(
                result.get("rfq_requirements", []),
                result.get("cad_findings", []),
                cad_bytes,
                cad_mime,
            )
        except Exception as annotation_exc:
            logger.warning(
                "[COMPARE] Auto-annotation failed: %s",
                annotation_exc,
                exc_info=True,
            )

        response_data = {
            "success": True,
            **result,
            "annotated_image": annotated_image,
            "annotations": annotation_records,
        }

        logger.info(
            "[COMPARE] Comparison result: match=%s, confidence=%s",
            result.get("match"),
            result.get("confidence"),
        )

        return JSONResponse(response_data)

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("[COMPARE] Unexpected error: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error comparing files: {str(exc)}")


@app.post("/compare-vendor")
async def compare_vendor_rfqs(
    files: TypingList[UploadFile] = File(...),
):
    """Upload multiple vendor RFQ documents and compare them."""
    try:
        allowed_rfq_types = {
            "application/pdf",
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }

        if not files or len(files) < 2:
            raise HTTPException(
                status_code=400,
                detail="At least 2 RFQ files are required for comparison",
            )

        # Validate all files
        rfq_inputs: List[Tuple[Union[bytes, str], Optional[str], str]] = []
        for idx, file in enumerate(files):
            if not file.content_type or file.content_type not in allowed_rfq_types:
                raise HTTPException(
                    status_code=400,
                    detail=f"File {idx + 1} ({file.filename}) must be a PDF or Word document (.doc, .docx)",
                )
            
            file_bytes = await file.read()
            if len(file_bytes) == 0:
                raise HTTPException(status_code=400, detail=f"File {idx + 1} ({file.filename}) is empty")

            max_size = 20 * 1024 * 1024
            if len(file_bytes) > max_size:
                raise HTTPException(status_code=400, detail=f"File {idx + 1} ({file.filename}) exceeds 20MB limit")

            rfq_mime = file.content_type if file.content_type in allowed_rfq_types else "application/pdf"
            rfq_input = inspector._prepare_rfq_input(file_bytes, rfq_mime)
            rfq_inputs.append((rfq_input[0], rfq_input[1], file.filename or f"file_{idx + 1}"))

        logger.info(
            "[VENDOR-COMPARE] Processing %d vendor RFQ files",
            len(rfq_inputs),
        )

        vendor_prompt = (
            "You are given multiple RFQ (Request for Quotation) documents from different vendors. "
            "All RFQs refer to the same item or item category, but each vendor may present information "
            "in different formats, wording styles, and layout structures.\n\n"

            "YOUR OBJECTIVE:\n"
            "Extract structured information from each RFQ and produce a normalized multi-vendor comparison "
            "strictly based on the content found inside the documents.\n\n"

            "GLOBAL RULES (APPLY TO ALL STEPS):\n"
            "1. Treat each RFQ as a separate vendor entry, even if vendor names look similar.\n"
            "2. Only extract information that explicitly appears in each RFQ.\n"
            "3. Never assume, guess, infer, or hallucinate missing data.\n"
            "4. If a field is missing for a vendor, set it to null.\n"
            "5. Never copy values from one vendor into another vendor’s fields.\n"
            "6. Never use external product knowledge of any kind.\n"
            "7. If numbers appear in words (e.g., 'five'), convert them into numeric form.\n"
            "8. Normalize currency into INR (₹) whenever possible.\n"
            "9. Use integers for day-based fields and floats for price fields.\n"
            "10. This prompt must work for ANY product category (automotive, electrical, industrial, etc.). "
            "    Never restrict assumptions to a specific domain.\n\n"

            "FIELDS TO EXTRACT (IF PRESENT):\n"
            "- certification_level\n"
            "- unit_price_inr\n"
            "- extended_price\n"
            "- quantity_discount\n"
            "- shipping_terms\n"
            "- delivery_initial_days\n"
            "- delivery_subsequent_days\n"
            "- delivery_emergency_days\n"
            "- warranty_period\n"
            "- technical.product_type\n"
            "- technical.part_number\n"
            "- technical.dimensions\n"
            "- technical.specifications\n\n"

            "OUTPUT FORMAT (STRICT JSON ONLY):\n"
            "{\n"
            "  \"vendors\": [\n"
            "    {\n"
            "      \"vendor_name\": \"\",\n"
            "      \"certification_level\": \"\",\n"
            "      \"pricing\": {\n"
            "        \"unit_price_inr\": 0,\n"
            "        \"extended_price\": 0,\n"
            "        \"quantity_discount\": \"\",\n"
            "        \"shipping_terms\": \"\"\n"
            "      },\n"
            "      \"delivery\": {\n"
            "        \"initial_days\": 0,\n"
            "        \"subsequent_days\": 0,\n"
            "        \"emergency_days\": 0\n"
            "      },\n"
            "      \"warranty\": \"\",\n"
            "      \"technical\": {\n"
            "        \"product_type\": \"\",\n"
            "        \"part_number\": \"\",\n"
            "        \"dimensions\": {},\n"
            "        \"specifications\": {}\n"
            "      }\n"
            "    }\n"
            "  ],\n"
            "  \"comparison\": {\n"
            "    \"best_price_vendor\": \"\",\n"
            "    \"best_delivery_vendor\": \"\",\n"
            "    \"best_warranty_vendor\": \"\",\n"
            "    \"overall_recommendation\": \"\"\n"
            "  }\n"
            "}\n\n"

            "RECOMMENDATION RULES (CRITICAL):\n"
            "- The 'overall_recommendation' field must recommend ONLY ONE vendor (the single best overall choice).\n"
            "- Select the vendor that offers the best overall value considering price, delivery time, warranty, and other factors.\n"
            "- The recommendation text must start with the vendor name and explain why this ONE vendor is recommended.\n"
            "- Do NOT recommend multiple vendors. Choose only the top 1 vendor.\n"
            "- Example format: \"[Vendor Name] offers the best overall value with [reasons].\"\n"
            "- The vendor name in overall_recommendation must exactly match one of the vendor_name values from the vendors array.\n\n"

            "FINAL OUTPUT RULES:\n"
            "- Output must be valid JSON only.\n"
            "- No markdown, no explanation, no commentary.\n"
            "- No additional notes before or after the JSON.\n"
            "- Do not hallucinate any fields or values.\n"
            "- Only include what the RFQs explicitly provide.\n"
            "- overall_recommendation must mention exactly ONE vendor name.\n"
        )

        # Prepare files for Gemini
        gemini_files: List[Tuple[Union[bytes, str], Optional[str]]] = []
        for rfq_input, mime_type, filename in rfq_inputs:
            gemini_files.append((rfq_input, mime_type))

        logger.info("[VENDOR-COMPARE] Sending %d files to Gemini", len(gemini_files))

        response_text = inspector.client.chat_with_files(vendor_prompt, gemini_files)

        logger.info("[VENDOR-COMPARE] Received response from Gemini")

        result = inspector.parse_vendor_comparison_response(response_text)

        if result is None:
            logger.error("[VENDOR-COMPARE] Failed to parse vendor comparison response")
            logger.error("[VENDOR-COMPARE] Raw response (first 2000 chars): %s", response_text[:2000] if response_text else "None")
            raise HTTPException(status_code=500, detail="Unable to parse vendor comparison response")

        logger.info(
            "[VENDOR-COMPARE] Parsed result: vendors=%d",
            len(result.get("vendors", [])),
        )

        response_data = {
            "success": True,
            **result,
        }

        return JSONResponse(response_data)

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("[VENDOR-COMPARE] Unexpected error: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error comparing vendor RFQs: {str(exc)}")

